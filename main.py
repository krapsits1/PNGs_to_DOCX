import os
from docx import Document
import pytesseract
from PIL import Image

# Specify the folder containing the .png files
input_folder = "vt yo"
output_docx = "extracted_text.docx"

# Create a new Word document
doc = Document()
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# Ensure the folder exists
if not os.path.exists(input_folder):
    print(f"The folder '{input_folder}' does not exist.")
    exit()

# Iterate over all .png files in the folder
for filename in os.listdir(input_folder):
    if filename.lower().endswith(".png"):
        file_path = os.path.join(input_folder, filename)

        try:
            # Open the image file
            image = Image.open(file_path)

            # Extract text from the image using pytesseract
            # lang='lav' specifies the language as Latvian
            extracted_text = pytesseract.image_to_string(image, lang='lav')

            # Add a heading for the image file name
            doc.add_heading(filename, level=2)

            # Add the extracted text to the document
            doc.add_paragraph(extracted_text)

        except Exception as e:
            print(f"Could not process file '{filename}'. Error: {e}")

# Save the extracted text to a .docx file
try:
    doc.save(output_docx)
    print(f"Text successfully extracted and saved to '{output_docx}'.")
except Exception as e:
    print(f"Could not save the Word document. Error: {e}")
